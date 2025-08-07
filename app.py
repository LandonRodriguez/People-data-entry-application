import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document

# Configure page
st.set_page_config(
    page_title="People Information Manager",
    page_icon="👥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        margin-bottom: 2rem;
    }
    
    .stats-container {
        display: flex;
        justify-content: space-around;
        margin: 1rem 0;
    }
    
    .stat-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        min-width: 120px;
    }
    
    .stat-number {
        font-size: 2em;
        font-weight: bold;
    }
    
    .stat-label {
        font-size: 0.9em;
        opacity: 0.9;
    }
    
    .person-card {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #667eea;
        margin: 0.5rem 0;
    }
    
    .success-message {
        background: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 5px;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    
    .error-message {
        background: #f8d7da;
        color: #721c24;
        padding: 1rem;
        border-radius: 5px;
        border-left: 4px solid #dc3545;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

class StreamlitPeopleManager:
    def __init__(self):
        # Initialize session state
        if 'people_data' not in st.session_state:
            st.session_state.people_data = []
        
        if 'success_message' not in st.session_state:
            st.session_state.success_message = ""
        
        if 'error_message' not in st.session_state:
            st.session_state.error_message = ""
    
    def add_person(self, first_name, last_name, age, job_title, city, state):
        """Add a person to the data"""
        person = {
            'First Name': first_name.strip(),
            'Last Name': last_name.strip(),
            'Age': int(age),
            'Job Title': job_title.strip(),
            'City': city.strip(),
            'State': state.strip()
        }
        
        st.session_state.people_data.append(person)
        st.session_state.success_message = f"✅ Added {first_name} {last_name} successfully!"
        st.session_state.error_message = ""
        return True
    
    def validate_input(self, first_name, last_name, age, job_title, city, state):
        """Validate form input"""
        if not all([first_name.strip(), last_name.strip(), job_title.strip(), city.strip(), state.strip()]):
            return False, "Please fill in all fields!"
        
        if age < 1 or age > 120:
            return False, "Please enter a valid age (1-120)!"
        
        return True, ""
    
    def get_statistics(self):
        """Calculate statistics from the data"""
        if not st.session_state.people_data:
            return 0, 0, 0
        
        total_count = len(st.session_state.people_data)
        avg_age = sum(person['Age'] for person in st.session_state.people_data) / total_count
        unique_states = len(set(person['State'] for person in st.session_state.people_data))
        
        return total_count, round(avg_age, 1), unique_states
    
    def create_excel_file(self):
        """Create Excel file in memory"""
        if not st.session_state.people_data:
            return None
        
        df = pd.DataFrame(st.session_state.people_data)
        
        # Create BytesIO buffer
        buffer = io.BytesIO()
        
        # Write Excel file to buffer
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='People Data', index=False)
            
            # Auto-adjust columns width
            worksheet = writer.sheets['People Data']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        buffer.seek(0)
        return buffer
    
    def create_word_file(self):
        """Create Word document in memory"""
        if not st.session_state.people_data:
            return None
        
        doc = Document()
        doc.add_heading('People Directory', 0)
        
        for person in st.session_state.people_data:
            sentence = f"{person['First Name']} {person['Last Name']}, {person['Age']} years old, works as a {person['Job Title']} and lives in {person['City']}, {person['State']}."
            doc.add_paragraph(sentence)
            doc.add_paragraph()  # Add spacing
        
        # Save to BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def clear_messages(self):
        """Clear success and error messages"""
        st.session_state.success_message = ""
        st.session_state.error_message = ""
    
    def display_people_list(self):
        """Display the list of people"""
        if not st.session_state.people_data:
            st.info("📝 No people added yet. Use the form to add your first person!")
            return
        
        for i, person in enumerate(st.session_state.people_data):
            st.markdown(f"""
            <div class="person-card">
                <strong>{person['First Name']} {person['Last Name']}</strong><br>
                Age: {person['Age']} • {person['Job Title']}<br>
                📍 {person['City']}, {person['State']}
            </div>
            """, unsafe_allow_html=True)

def main():
    # Initialize the manager
    manager = StreamlitPeopleManager()
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>👥 People Information Manager</h1>
        <p>Add people's information and export to Excel/Word with ease</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Layout with columns
    col1, col2 = st.columns([1, 1])
    
    # Left column - Form
    with col1:
        st.subheader("➕ Add New Person")
        
        with st.form("person_form", clear_on_submit=True):
            first_name = st.text_input("First Name", placeholder="Enter first name")
            last_name = st.text_input("Last Name", placeholder="Enter last name")
            age = st.number_input("Age", min_value=1, max_value=120, value=25)
            job_title = st.text_input("Job Title", placeholder="e.g., Software Engineer")
            city = st.text_input("City", placeholder="Enter city")
            state = st.text_input("State", placeholder="Enter state")
            
            submitted = st.form_submit_button("Add Person", type="primary", use_container_width=True)
            
            if submitted:
                manager.clear_messages()
                
                # Validate input
                is_valid, error_msg = manager.validate_input(first_name, last_name, age, job_title, city, state)
                
                if is_valid:
                    manager.add_person(first_name, last_name, age, job_title, city, state)
                    st.rerun()
                else:
                    st.session_state.error_message = error_msg
                    st.rerun()
    
    # Right column - Statistics and data
    with col2:
        st.subheader("📊 Statistics")
        
        # Get statistics
        total_count, avg_age, unique_states = manager.get_statistics()
        
        # Display statistics
        col_stat1, col_stat2, col_stat3 = st.columns(3)
        
        with col_stat1:
            st.metric("Total People", total_count)
        
        with col_stat2:
            st.metric("Average Age", f"{avg_age}" if total_count > 0 else "0")
        
        with col_stat3:
            st.metric("Unique States", unique_states)
    
    # Messages
    if st.session_state.success_message:
        st.success(st.session_state.success_message)
    
    if st.session_state.error_message:
        st.error(st.session_state.error_message)
    
    # People directory
    st.subheader("📋 People Directory")
    manager.display_people_list()
    
    # Action buttons
    if st.session_state.people_data:
        st.subheader("📥 Export Options")
        
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        with col_btn1:
            # Excel download
            excel_buffer = manager.create_excel_file()
            if excel_buffer:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"people_data_{timestamp}.xlsx"
                
                st.download_button(
                    label="📊 Download Excel",
                    data=excel_buffer.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary",
                    use_container_width=True
                )
        
        with col_btn2:
            # Word download
            word_buffer = manager.create_word_file()
            if word_buffer:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"people_profiles_{timestamp}.docx"
                
                st.download_button(
                    label="📄 Download Word",
                    data=word_buffer.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col_btn3:
            # Clear all data
            if st.button("🗑️ Clear All Data", type="secondary", use_container_width=True):
                if st.session_state.get('confirm_clear', False):
                    st.session_state.people_data = []
                    st.session_state.success_message = "All data cleared successfully!"
                    st.session_state.confirm_clear = False
                    st.rerun()
                else:
                    st.session_state.confirm_clear = True
                    st.warning("⚠️ Click again to confirm deletion of all data!")
                    st.rerun()
    
    # Sidebar with additional info
    with st.sidebar:
        st.header("ℹ️ About")
        st.info("""
        **People Information Manager**
        
        This application helps you:
        - Add and manage people's information
        - View real-time statistics
        - Export data to Excel or Word formats
        - Keep everything organized
        
        **How to use:**
        1. Fill the form with person's details
        2. Click 'Add Person'
        3. Repeat for more people
        4. Export when ready
        """)
        
        if st.session_state.people_data:
            st.header("📈 Data Overview")
            df = pd.DataFrame(st.session_state.people_data)
            
            # Age distribution
            st.subheader("Age Distribution")
            st.bar_chart(df['Age'].value_counts().sort_index())
            
            # State distribution
            if len(df['State'].unique()) > 1:
                st.subheader("People by State")
                state_counts = df['State'].value_counts()
                st.bar_chart(state_counts)

if __name__ == "__main__":
    main()
